"""
build_explainer_docx.py
-----------------------
Render docs/PAPER_EXPLAINED.md as a study-guide .docx in the house style used
for Busra's viva guides: Calibri, blue/navy headings, tinted single-cell boxes,
navy-header comparison tables.

The guide is bilingual. English blocks are set plain; Turkish blocks are set in
a tinted box so the two languages stay visually separable while sitting next to
each other on the page.

Usage:
    python scripts/build_explainer_docx.py

Output:
    docs/PAPER_EXPLAINED.docx
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).parent.parent
SRC = ROOT / "docs" / "PAPER_EXPLAINED.md"
OUT = ROOT / "docs" / "PAPER_EXPLAINED.docx"

BLUE = RGBColor(0x0F, 0x4C, 0x81)
NAVY = RGBColor(0x16, 0x21, 0x3E)
GREY = RGBColor(0x44, 0x44, 0x44)
BODY = RGBColor(0x33, 0x33, 0x33)
TR_TEXT = RGBColor(0x1B, 0x4D, 0x3E)

FILL_TR = "EEF6F2"      # Turkish blocks
FILL_KEY = "E8F1FA"     # "in one breath" style call-outs
NAVY_HEX = "16213E"


def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def strip_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        borders.append(e)
    table._tbl.tblPr.append(borders)


def inline(par, text, *, size=10.5, color=BODY, bold=False):
    """Write text into a paragraph, honouring **bold**, *italic* and `code`."""
    for part in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", text):
        if not part:
            continue
        run = par.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text, run.bold = part[2:-2], True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
        elif part.startswith("*") and part.endswith("*"):
            run.text, run.italic = part[1:-1], True
        else:
            run.text = part
            run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        if run.font.name != "Consolas":
            run.font.name = "Calibri"


def add_box(doc, lines, fill, text_color):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    strip_borders(t)
    cell = t.cell(0, 0)
    shade(cell, fill)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        inline(p, line, color=text_color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        shade(c, NAVY_HEX)
        c.paragraphs[0].text = ""
        inline(c.paragraphs[0], h, size=9.5, color=RGBColor(0xFF, 0xFF, 0xFF),
               bold=True)
    for r in body:
        cells = t.add_row().cells
        for i, v in enumerate(r[:len(header)]):
            cells[i].paragraphs[0].text = ""
            inline(cells[i].paragraphs[0], v, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines, i):
    """Collect a markdown pipe-table starting at lines[i]. Returns rows, next_i."""
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def main():
    md = SRC.read_text().split("\n")
    doc = Document()

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.font.color.rgb = BODY
    st.paragraph_format.space_after = Pt(6)

    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.1)
        s.top_margin = s.bottom_margin = Cm(1.9)
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("dMRI Rosetta Stone — Study Companion / Çalışma Rehberi")
        r.font.size = Pt(8)
        r.font.color.rgb = GREY

    i, lang = 0, None
    while i < len(md):
        line = md[i].rstrip()

        if not line:
            i += 1
            continue

        if line.startswith("|"):
            rows, i = parse_table(md, i)
            if rows:
                add_table(doc, rows)
            continue

        if line.startswith("---"):
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(24)
            r.font.color.rgb = BLUE
            p.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_page_break()
            p = doc.add_paragraph()
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(16)
            r.font.color.rgb = NAVY
            p.paragraph_format.space_after = Pt(8)
            lang = None
            i += 1
            continue

        if line.startswith("### "):
            title = line[4:]
            if title in ("EN", "TR"):
                lang = title
                i += 1
                continue
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(12.5)
            r.font.color.rgb = BLUE
            p.paragraph_format.space_before = Pt(8)
            lang = None
            i += 1
            continue

        # Collect a run of body lines belonging to the same block.
        block = []
        while i < len(md) and md[i].strip() and not md[i].lstrip().startswith(
                ("#", "|", "---")):
            block.append(md[i].strip())
            i += 1

        text = " ".join(block)

        # Language is signalled two ways: by an "### EN" / "### TR" heading,
        # or inline as "**TR — Ne.**". The inline form has to persist across
        # the paragraphs that follow it, otherwise a Turkish "**Neden.**"
        # paragraph is set as English and loses its tint.
        if re.match(r"\*\*?(TR|EN)\b", text):
            lang = "TR" if text.lstrip("*").startswith("TR") else "EN"
        turkish = lang == "TR"

        # A few paired paragraphs carry no marker at all. Fall back on letters
        # that exist in Turkish but not in English.
        if not turkish and lang is None and re.search(r"[ışğİŞĞ]", text):
            turkish = True

        if turkish:
            add_box(doc, [text], FILL_TR, TR_TEXT)
        elif text.startswith(("**Q:", "**Point", "**Honest", "**Kullanılacak")):
            add_box(doc, [text], FILL_KEY, BLUE)
        else:
            p = doc.add_paragraph()
            if text.startswith(("- ", "* ", "1. ", "2. ", "3. ", "4. ",
                                "5. ", "6. ")):
                p.paragraph_format.left_indent = Cm(0.6)
            inline(p, text)

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
